#!/usr/bin/env python3
"""Create a trimmed playtester rulebook PDF for The Quiet Vale."""

from __future__ import annotations

import argparse
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from pypdf import PdfReader
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from export_styled_rulebook_from_docx import (
    DEFAULT_SOURCE,
    OUT_DIR,
    PALETTE,
    STYLES,
    clean,
    docx_table_to_flowable,
    prepared_cover_logo,
    safe,
)


RULEBOOK_VERSION = "v3.1 current build"
CURRENT_BUILD_SOURCE = (
    Path(__file__).resolve().parents[1]
    / "exports"
    / "current_build_documents"
    / "The_Quiet_Vale_Player_Facing_Rulebook_v3_1_Current_Build.docx"
)
OUTPUT = OUT_DIR / "The_Quiet_Vale_Playtester_Rulebook_Current_Build_v3_1.pdf"
MANIFEST = OUT_DIR / "manifest_playtester_current_build_v3_1.txt"


PLAYTEST_STYLES = {
    **STYLES,
    "playtest_cover_note": ParagraphStyle(
        "PlaytestCoverNote",
        parent=STYLES["cover_small"],
        fontName="Helvetica",
        fontSize=7.4,
        leading=10,
        textColor=PALETTE["brass"],
        alignment=TA_CENTER,
        spaceAfter=4,
    ),
    "flavour": ParagraphStyle(
        "QVFlavour",
        parent=STYLES["body"],
        fontName="Times-Italic",
        fontSize=10.4,
        leading=14.4,
        textColor=PALETTE["graphite"],
        spaceAfter=7,
    ),
}


def safe_flavour(value: str) -> str:
    """Preserve flavour wording and punctuation while escaping PDF markup."""
    return escape(" ".join(str(value or "").split()))


def heading(text: str, level: int = 1) -> Paragraph:
    return Paragraph(safe(text), PLAYTEST_STYLES["h1" if level == 1 else "h2"])


def body(text: str) -> Paragraph:
    return Paragraph(safe(text), PLAYTEST_STYLES["body"])


def bullet(text: str) -> Paragraph:
    return Paragraph(f'<font color="#7B5E2E">-</font> {safe(text)}', PLAYTEST_STYLES["body"])


def numbered(items: list[str]) -> list[Paragraph]:
    return [Paragraph(f'<font color="#7B5E2E"><b>{index}.</b></font> {safe(text)}', PLAYTEST_STYLES["body"]) for index, text in enumerate(items, start=1)]


def cover_story() -> list:
    story = [Spacer(1, 50 * mm)]
    logo = prepared_cover_logo()
    if logo:
        image = Image(str(logo), width=18 * mm, height=22 * mm)
        image.hAlign = "CENTER"
        story.extend([image, Spacer(1, 16 * mm)])

    story.extend(
        [
            Paragraph("The Quiet Vale", PLAYTEST_STYLES["cover_title"]),
            Paragraph("Seasons of Settlement", PLAYTEST_STYLES["cover_subtitle"]),
            Paragraph("Playtester Rulebook", PLAYTEST_STYLES["playtest_cover_note"]),
            Spacer(1, 56 * mm),
            Paragraph("A focused table reference for learning and blind playtesting.", PLAYTEST_STYLES["cover_body"]),
            PageBreak(),
        ]
    )
    return story


def table_by_index(doc, index: int):
    return doc.tables[index]


def front_flavour(doc) -> list[str]:
    lines = []
    found_title_block = False
    for paragraph in doc.paragraphs:
        text = " ".join((paragraph.text or "").split())
        if not text:
            continue
        if text == "Contents":
            break
        if text == "The banners of war have long since fallen silent.":
            found_title_block = True
        if found_title_block:
            lines.append(text)
    return lines


def compact_rule_box(label: str, text: str):
    table = Table(
        [[Paragraph(safe(label.upper()), PLAYTEST_STYLES["note_label"]), Paragraph(safe(text), PLAYTEST_STYLES["note_body"])]],
        colWidths=[35 * mm, 124 * mm],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), PALETTE["note"]),
                ("BOX", (0, 0), (-1, -1), 0.65, PALETTE["line"]),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def manual_table(rows: list[list[str]], widths: list[float] | None = None):
    data = []
    for row_index, row in enumerate(rows):
        style = PLAYTEST_STYLES["table_header"] if row_index == 0 else PLAYTEST_STYLES["table_cell"]
        data.append([Paragraph(safe(cell), style) for cell in row])

    flowable = Table(data, colWidths=widths or None, hAlign="LEFT", repeatRows=1)
    flowable.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), PALETTE["graphite"]),
                ("TEXTCOLOR", (0, 0), (-1, 0), PALETTE["parchment_light"]),
                ("BACKGROUND", (0, 1), (-1, -1), PALETTE["parchment_light"]),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [PALETTE["parchment_light"], PALETTE["parchment_deep"]]),
                ("GRID", (0, 0), (-1, -1), 0.32, PALETTE["line"]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3.6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3.8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return flowable


def steward_house_setup_table():
    return manual_table(
        [
            ["Steward", "Free setup token placement", "First normal tile"],
            ["Vanguard", "Steward token on Woodland.", "Player choice."],
            ["Knight", "Steward token on Arable Land.", "Player choice."],
            ["Sentinel", "Steward token on Mountains.", "Player choice."],
            ["Ranger", "Steward token on Heaths.", "Player choice."],
            ["Warden", "Steward token on Ruins.", "Player choice."],
            [
                "Quartermaster",
                "Steward token on Woodland, Mountains, Heaths, Arable Land, or Ruins. Not Grasslands or River.",
                "Player choice.",
            ],
        ],
        widths=[28 * mm, 86 * mm, 45 * mm],
    )


def online_encounter_setup_table():
    return manual_table(
        [
            ["Players", "Boons", "Burdens", "Arrivals", "Golden Boons", "Hidden Cards", "Deck Cards"],
            ["1", "5", "5", "5", "0 - not supported online", "9", "3 standard"],
            ["2", "10", "10", "10", "0 - not supported online", "9 each", "6 standard"],
            ["3", "15", "15", "15", "0 - not supported online", "9 each", "9 standard"],
            ["4", "20", "20", "20", "0 - not supported online", "9 each", "12 standard"],
        ],
        widths=[15 * mm, 17 * mm, 19 * mm, 19 * mm, 42 * mm, 23 * mm, 24 * mm],
    )


def online_round_phase_table():
    return manual_table(
        [
            ["Phase", "Name", "Summary"],
            ["1", "Seasonal Seed Encounter Cards", "Only in Rounds 1, 5, and 9. Each player seeds 3 hidden cards: 1 Top, 1 Middle, and 1 Bottom."],
            ["2", "Reveal Encounters", "Reveal standard Encounter Cards equal to player count. Golden Boons are not used in the online prototype."],
            ["3", "Player Turns", "Each player takes one turn with 4 Actions. The group chooses turn order."],
            ["4", "End of Round", "Remove Arrival timers, expire failed Arrivals, discard applicable Boons, and advance the Round Timer."],
        ],
        widths=[16 * mm, 42 * mm, 101 * mm],
    )


def online_encounter_type_table():
    return manual_table(
        [
            ["Encounter Type", "When Revealed"],
            ["Boon", "Resolve the current Season effect. If it modifies a later action or future event, keep it face-up according to its text."],
            ["Burden", "Apply the current Season effect. Place the card on the Stewards Board as an active Burden."],
            ["Arrival", "Place it on the Stewards Board with 3 timer tokens. Players may complete it by fulfilling its Requirement and spending an Interact action."],
            ["Golden Boon", "Not currently supported by the online prototype."],
        ],
        widths=[36 * mm, 123 * mm],
    )


def steward_power_table():
    return manual_table(
        [
            ["Steward", "Once per Season Steward Power"],
            [
                "Vanguard",
                "Place a Travel Tile or Resource Tile without spending the Place Action. Costs and requirements still apply.",
            ],
            [
                "Knight",
                "Place a Housing Tile without spending the Place Action. Costs and requirements still apply.",
            ],
            [
                "Sentinel",
                "Upgrade a Core Tile without spending the Upgrade Action. Costs and requirements still apply.",
            ],
            [
                "Ranger",
                "Once per Season, travel to anywhere for free before taking a map action. This does not spend an Action.",
            ],
            [
                "Quartermaster",
                "Exchange up to 3 Warehouse resources for the same number of different non-Goods resources.",
            ],
            [
                "Warden",
                "Resolve an active Burden without spending the Resolve Action. Normal costs and requirements still apply.",
            ],
        ],
        widths=[43 * mm, 116 * mm],
    )


def draw_playtester_cover(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(PALETTE["smoke"])
    canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
    canvas.setStrokeColor(PALETTE["brass_dark"])
    canvas.setLineWidth(1.0)
    canvas.rect(12 * mm, 12 * mm, A4[0] - 24 * mm, A4[1] - 24 * mm, fill=0, stroke=1)
    canvas.setStrokeColor(PALETTE["brass"])
    canvas.setLineWidth(0.45)
    canvas.rect(16 * mm, 16 * mm, A4[0] - 32 * mm, A4[1] - 32 * mm, fill=0, stroke=1)
    canvas.line(30 * mm, A4[1] - 83 * mm, 88 * mm, A4[1] - 83 * mm)
    canvas.line(122 * mm, A4[1] - 83 * mm, A4[0] - 30 * mm, A4[1] - 83 * mm)
    canvas.line(30 * mm, 74 * mm, A4[0] - 30 * mm, 74 * mm)
    canvas.setFillColor(PALETTE["brass"])
    canvas.setFont("Helvetica", 6.5)
    canvas.drawCentredString(A4[0] / 2, 34 * mm, "Copyright (C) 2026 Robert Little. All rights reserved.")
    canvas.drawCentredString(A4[0] / 2, 24 * mm, "www.thequietvalegame.com")
    canvas.restoreState()


def draw_playtester_body(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(PALETTE["parchment"])
    canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
    canvas.setStrokeColor(PALETTE["line"])
    canvas.setLineWidth(0.65)
    canvas.rect(13 * mm, 12 * mm, A4[0] - 26 * mm, A4[1] - 24 * mm, fill=0, stroke=1)
    canvas.setStrokeColor(PALETTE["brass"])
    canvas.setLineWidth(0.35)
    canvas.line(20 * mm, A4[1] - 17 * mm, A4[0] - 20 * mm, A4[1] - 17 * mm)
    canvas.line(20 * mm, 17 * mm, A4[0] - 20 * mm, 17 * mm)
    canvas.setFont("Helvetica", 7.2)
    canvas.setFillColor(PALETTE["muted"])
    canvas.drawString(20 * mm, 10 * mm, "The Quiet Vale: Seasons of Settlement")
    canvas.drawRightString(A4[0] - 20 * mm, 10 * mm, f"Playtester draft {RULEBOOK_VERSION} - Page {doc.page}")
    canvas.restoreState()


def build_story(doc) -> list:
    story = cover_story()

    story.append(heading("Welcome To The Vale"))
    for line in front_flavour(doc):
        story.append(Paragraph(safe_flavour(line), PLAYTEST_STYLES["flavour"]))
    story.append(Spacer(1, 4))

    story.append(heading("How To Use This Playtest Guide"))
    story.append(body("This is a trimmed table reference for learning the game. It keeps the core flow, costs, timing, and card handling rules players need during a playtest. The fuller production rulebook can keep the deeper wording and edge-case lists."))
    story.append(compact_rule_box("Use card text", "When a card gives a specific exception, use the card. This guide explains the normal rules and avoids listing every card-specific fallback."))

    story.append(heading("Game Overview"))
    story.append(body("The Quiet Vale is a cooperative tile-laying and settlement-building game for 1-4 players. The group builds a shared settlement across three Seasons, manages a shared Warehouse, reveals Encounters, welcomes Arrivals, and prevents Strain from disabling important tiles."))
    for item in [
        "Build and upgrade a shared settlement.",
        "Gather and spend shared resources from the Warehouse.",
        "Seed hidden Encounter Cards at the start of each Season.",
        "Complete Arrivals to unlock Special Tiles.",
        "Resolve or endure Burdens before Strain becomes too costly.",
    ]:
        story.append(bullet(item))

    story.append(heading("Setup"))
    story.extend(
        numbered(
            [
                "Lay out the Redesigned Basic Map v0.2, Stewards Board, Warehouse Board, and Round Timer.",
                "Each player chooses a unique Steward and takes that Steward's Player Aid and Steward token.",
                "Set the shared Warehouse using the player-count table.",
                "Build the balanced Encounter pool, deal 9 hidden Encounter Cards to each player, and deal 3 standard Encounter Cards per player to the Encounter Deck. Golden Boons are not currently supported by the online prototype.",
                "Each player places their Steward token for free on its setup terrain. This costs 0 Actions and 0 resources, ignores normal adjacency, and must use an empty non-River hex.",
                "Begin Round 1 with Seasonal Seed Encounter Cards. Each player seeds one card to the Top, one to the Middle, and one to the Bottom. There is no forced opening Resource tile; players choose their first normal tile action.",
            ]
        )
    )
    story.append(heading("Starting Warehouse", 2))
    story.append(docx_table_to_flowable(table_by_index(doc, 0)))
    story.append(compact_rule_box("Locked map", "Online playtests use Redesigned Basic Map v0.2. Older map options are retained only as repository reference files and are not selectable during setup."))
    story.append(heading("Steward Token Setup", 2))
    story.append(steward_house_setup_table())
    story.append(compact_rule_box("Steward powers", "Steward House tiles are not used in this playtest version. Each Steward has their once-per-Season power from the start of the game. Only that Steward may use their own power."))
    story.append(heading("Encounter Setup", 2))
    story.append(online_encounter_setup_table())
    story.append(compact_rule_box("Online prototype note", "Golden Boons are excluded from the current online prototype and should not be added during online playtests."))

    story.append(heading("Round Structure"))
    story.append(body("The game lasts 12 rounds across three Seasons. Rounds 1, 5, and 9 begin with Seasonal Seed Encounter Cards. All other rounds skip seeding and begin with Reveal Encounters."))
    story.append(
        manual_table(
            [
                ["Season", "Rounds"],
                ["I", "Rounds 1-4"],
                ["II", "Rounds 5-8"],
                ["III", "Rounds 9-12"],
            ],
            widths=[40 * mm, 119 * mm],
        )
    )
    story.append(online_round_phase_table())
    story.append(compact_rule_box("Seasonal seeding", "At the start of each Season, each player seeds 3 hidden cards: one Top, one Middle, and one Bottom. Players do not seed cards during the other rounds."))
    story.append(heading("Reveal And Actions By Player Count", 2))
    story.append(docx_table_to_flowable(table_by_index(doc, 6)))

    story.append(heading("Player Actions"))
    story.append(docx_table_to_flowable(table_by_index(doc, 7)))
    story.append(body("There is no normal disconnected Travel action in this playtest version. A Steward must place, upgrade, and activate tiles on their own connected settlement network unless a Steward power or card explicitly says otherwise."))

    story.append(heading("Map, Tiles, And Reachability"))
    for item in [
        "Tiles stay on the map unless a rule explicitly removes them.",
        "Placed, non-Overstrained tiles connected by side adjacency form the connected settlement network.",
        "Stewards share the connected settlement network; one player's placed tiles can extend another player's reach if connected.",
        "Each player's Steward location starts at their setup token and later becomes the tile they last interacted with or placed.",
        "A tile is reachable if it is the Steward's tile or connected to that tile through the connected settlement network.",
        "Overstrained tiles break the connected settlement network.",
        "Multi-hex tiles count all hexes in their footprint for adjacency and connection. Single-hex tiles do not need rotation.",
    ]:
        story.append(bullet(item))
    story.append(heading("Rivers And Bridges", 2))
    for item in [
        "River hexes are barriers. Tiles do not connect across a river unless a Bridge or card effect allows the crossing.",
        "Every River hex is a legal potential Bridge placement site.",
        "Without a Bridge or explicit crossing rule, a river break means the far side is not connected for normal Steward actions.",
        "A Bridge connects settlement networks through its river hex. An Overstrained Bridge stops providing that connection.",
    ]:
        story.append(bullet(item))

    story.append(
        KeepTogether(
            [
                heading("Resources And Warehouse"),
                body("The Warehouse is shared by all players. It stores Wood, Stone, Metal, Food, Herbs, and Goods. The Warehouse limit is 15 of each resource."),
                docx_table_to_flowable(table_by_index(doc, 8)),
            ]
        )
    )

    story.append(PageBreak())
    story.append(heading("Stewards And Powers"))
    for item in [
        "Each Steward token is placed for free during setup before the first Seasonal Seed Encounter Cards step.",
        "The acting Steward's marker starts on their own setup terrain after token placement.",
        "Steward House tiles are not used in this playtest version.",
        "Each Steward has their once-per-Season power from game start.",
        "Only the matching Steward may use their own power, even though all Stewards may use the shared connected settlement network.",
        "Using a Steward Power never waives resource costs or placement requirements unless that power says so.",
    ]:
        story.append(bullet(item))
    story.append(steward_power_table())

    story.append(heading("Encounter Cards"))
    story.append(body("Encounter Cards are seeded into and revealed from the Encounter Deck. Seasonal seeding happens only in Rounds 1, 5, and 9. Cards on the Stewards Board are open information. Players may inspect active cards, completed Arrivals, active Burdens, and face-up Boons."))
    story.append(online_encounter_type_table())

    story.append(heading("Arrivals And Special Tiles"))
    for item in [
        "An Arrival enters play with 3 timer tokens.",
        "To complete an Arrival, fulfill its Requirement all at once and spend 1 Action to interact with it.",
        "Arrival resource Requirements cannot be partly paid over time.",
        "At the end of each round, each active Arrival loses 1 timer token. If it reaches 0 before completion, it expires and the group adds 1 Strain to any placed tile of their choice.",
        "A completed Arrival unlocks its Special Tile. Any player may later place that unlocked Special Tile with a normal Place a Tile action.",
        "Unlocked but unplaced Special Tiles are not on the map and have no effects until placed.",
    ]:
        story.append(bullet(item))

    story.append(PageBreak())
    story.append(heading("Burdens"))
    for item in [
        "When a Burden is revealed, apply the current Season effect and then place it on the Stewards Board as an active Burden.",
        "Unresolved active Burdens reapply at the start of Round 5 using Season II text and at the start of Round 9 using Season III text.",
        "If a Burden places Strain, choose valid tiles with fewer than 3 Strain. If there are fewer valid tiles than requested, choose all valid tiles.",
        "If no valid target exists, no Strain is placed for that application. The Burden remains active unless resolved.",
        "Some Season III Burdens include fallback targets or Season III-only resolution. Use the text printed on that card rather than memorising a separate list.",
        "Resolving a Burden removes it from the Stewards Board and prevents future reapplications. It does not remove Strain already placed.",
    ]:
        story.append(bullet(item))
    story.append(compact_rule_box("Playtest focus", "Players should notice active Burdens. Leaving them unresolved can create repeated pressure and final score penalties."))

    story.append(PageBreak())
    story.append(heading("Boons And Golden Boons"))
    for item in [
        "When a standard Boon is revealed, apply the effect matching the current Season.",
        "Boon effects are optional unless the card says otherwise.",
        "Most Boons resolve immediately and are discarded.",
        "Boons that affect the next or first eligible action stay face-up until used, then are discarded. If unused, discard them in Phase Four unless the card says to keep it face-up.",
        "Golden Boons are not currently supported by the online prototype, so do not add them for online blind tests.",
    ]:
        story.append(bullet(item))

    story.append(heading("Strain, Supported, And Overstrained"))
    story.append(body("Strain represents pressure on the settlement: damage, exhaustion, unstable resources, and unresolved social harm."))
    for item in [
        "A tile cannot have more than 3 Strain.",
        "A tile with 3 Strain is Overstrained.",
        "Tiles with 1 or 2 Strain still function unless a rule says otherwise.",
        "Supported prevents the first Strain token applied to that tile each round.",
        "Supported does not stack. Multiple sources still prevent only the first Strain that round.",
        "Overstrained tiles cannot be activated or upgraded, provide no passive effects, do not score, and do not contribute to reachability.",
        "As soon as a tile has fewer than 3 Strain, it is no longer Overstrained.",
    ]:
        story.append(bullet(item))
    story.append(compact_rule_box("Support tile update", "Tiles that give broad adjacent Supported now do so by activation. Basic support tiles give Supported to 1 adjacent eligible tile. Upgraded support tiles give Supported to up to 2 adjacent eligible tiles."))
    story.append(compact_rule_box("Placement discount update", "Brewery of Legends and Labourers' Yard may apply their adjacent placement discount once per Season, not once per round."))

    story.append(heading("Season Ends And Final Scoring"))
    story.append(
        manual_table(
            [
                ["Timing", "Effect"],
                ["End of Season I, after Round 4", "Each Overstrained tile spreads 1 Strain to an adjacent tile, following normal Strain cap and prevention rules."],
                ["End of Season II, after Round 8", "Each Overstrained tile spreads 1 Strain to an adjacent tile, following normal Strain cap and prevention rules."],
                ["End of Season III, after Round 12", "Apply final Renown penalties: -6 Renown for each active Burden and -2 Renown for each Strain token on the board."],
            ],
            widths=[50 * mm, 109 * mm],
        )
    )
    for item in [
        "Add Population from non-Overstrained placed tiles.",
        "Add Renown from non-Overstrained placed tiles.",
        "Subtract 6 Renown for each unresolved active Burden.",
        "Subtract 2 Renown for each Strain token on the board.",
    ]:
        story.append(bullet(item))
    story.append(compact_rule_box("Final score", "Population + Renown - active Burden penalties - Strain penalties."))

    story.append(PageBreak())
    story.append(heading("Quick Glossary"))
    story.append(docx_table_to_flowable(table_by_index(doc, 12)))

    story.append(heading("Playtest Notes"))
    for item in [
        "Use this guide to learn and play. Use the cards themselves for exact one-off card rules.",
        "If a rule feels unclear, note what happened, the round, the card or tile involved, and what the group expected.",
        "The most useful feedback is about pacing, tension, choices, resource pressure, and whether the settlement story is coming through.",
    ]:
        story.append(bullet(item))
    story.append(Spacer(1, 8))
    story.append(
        compact_rule_box(
            "Ownership",
            "The Quiet Vale(TM): Seasons of Settlement and all prototype rulebook materials are (C) Robert Little. All rights reserved. Shared for private playtesting and review.",
        )
    )

    return story


def build_pdf(source: Path, output: Path):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(source)
    pdf = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=25 * mm,
        leftMargin=25 * mm,
        topMargin=24 * mm,
        bottomMargin=24 * mm,
        title=f"The Quiet Vale Playtester Rulebook Styled Draft {RULEBOOK_VERSION}",
        author="Robert Little",
        subject="Focused The Quiet Vale playtester rulebook draft",
    )
    pdf.build(build_story(doc), onFirstPage=draw_playtester_cover, onLaterPages=draw_playtester_body)
    pages = len(PdfReader(str(output)).pages)
    MANIFEST.write_text(
        "\n".join(
            [
                f"The Quiet Vale playtester rulebook styled draft {RULEBOOK_VERSION}",
                f"PDF: {output.name}",
                f"Source DOCX: {source}",
                f"Pages: {pages}",
                "Trim approach: table-facing rules only; opening flavour text preserved; current Steward token setup, locked v0.2 map, seasonal seeding, reachability, support, and scoring rules reflected.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(output)
    print(f"pages={pages}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=CURRENT_BUILD_SOURCE if CURRENT_BUILD_SOURCE.exists() else DEFAULT_SOURCE)
    parser.add_argument("--out", type=Path, default=OUTPUT)
    args = parser.parse_args()
    build_pdf(args.source, args.out)


if __name__ == "__main__":
    main()
